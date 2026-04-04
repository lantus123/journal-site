"""
Process .docx work manuals into chunked JSON for protocol gap analysis.

Usage:
    python scripts/process_manuals.py [--dept newborn] [--no-keywords]

Steps:
    1. Read all .docx files from config/manuals/
    2. Split by headings into semantic chunks
    3. Use Claude Haiku to extract keywords for each chunk
    4. Save to data/{dept}/manual_chunks.json
"""

import argparse
import json
import logging
import re
import sys
from pathlib import Path

from docx import Document

# Add project root to path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from src.llm import LLMClient

logging.basicConfig(level=logging.INFO, format="%(message)s")
logger = logging.getLogger(__name__)

KEYWORD_PROMPT = """你是新生兒科醫學專家。以下是一段科內工作手冊的內容。
請提取 5-15 個關鍵字，用於將來比對 PubMed 文獻時能匹配到這段 protocol。

要求：
- 包含中文和英文醫學術語（如「高頻震盪」和「HFOV」都要列）
- 包含藥物名、疾病名、治療方式、監測指標
- 包含常見縮寫（如 GA, BW, RDS, NEC, PDA 等）
- 不要太泛（如「新生兒」「治療」）

回傳 JSON array，例如：["HFOV", "高頻震盪", "MAP", "RDS", "呼吸窘迫症"]

手冊段落標題：{heading}
內容：
{content}
"""

# Minimum chunk size (characters) to process
MIN_CHUNK_SIZE = 50
# Maximum chunk size for LLM processing
MAX_CHUNK_SIZE = 8000


def _is_heading(para) -> tuple[bool, int]:
    """Detect headings by style name, font size, or bold formatting.

    Google Docs exports often lose Heading styles, so we also check:
    - Large font size (> 14pt / 177800 EMU) -> level 1
    - Medium font size (> 12pt but body-like) -> level 2
    - All-bold, short text (<=80 chars) -> level 2
    - Explicit Heading style -> use its level

    Returns (is_heading, level).
    """
    style = para.style.name if para.style else ""
    text = para.text.strip()
    runs = para.runs

    # 1. Explicit heading style
    if style.startswith("Heading"):
        try:
            level = int(style.replace("Heading", "").strip())
        except ValueError:
            level = 1
        return True, level

    if not runs or not text:
        return False, 0

    # 2. Large font -> chapter title (level 1)
    first_size = runs[0].font.size
    if first_size and first_size > 228600:  # > ~18pt
        return True, 1

    # 3. Medium-large font -> section heading (level 2)
    if first_size and first_size > 177800 and len(text) <= 80:  # > 14pt
        return True, 2

    # 4. All-bold, short text -> subsection heading (level 2)
    if len(text) <= 80:
        bold_runs = [r for r in runs if r.text.strip()]
        if bold_runs and all(r.bold for r in bold_runs):
            return True, 2

    return False, 0


def extract_chunks_from_docx(filepath: Path) -> list[dict]:
    """Extract text chunks split by headings from a .docx file."""
    doc = Document(filepath)
    # Derive chapter name from filename, e.g. "2025_NB_manual_03_Respiratory.docx" -> "Respiratory"
    stem = filepath.stem
    # Try to extract the topic part after the last underscore or number
    parts = stem.split("_")
    # Find the meaningful part (skip "2025", "NB", "manual", numbers)
    chapter = parts[-1] if len(parts) > 1 else stem

    chunks = []
    current_text = []
    heading_stack = [chapter]

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        is_head, level = _is_heading(para)

        if is_head:
            # Save previous chunk
            if current_text:
                content = "\n".join(current_text)
                if len(content) >= MIN_CHUNK_SIZE:
                    chunks.append({
                        "path": " / ".join(heading_stack),
                        "content": content,
                        "source_file": filepath.name,
                    })
                current_text = []

            # Update heading stack based on level
            heading_stack = heading_stack[:level]
            heading_stack.append(text)
        else:
            current_text.append(text)

    # Don't forget the last chunk
    if current_text:
        content = "\n".join(current_text)
        if len(content) >= MIN_CHUNK_SIZE:
            chunks.append({
                "path": " / ".join(heading_stack),
                "content": content,
                "source_file": filepath.name,
            })

    return chunks


def extract_chunks_from_docx_by_size(filepath: Path) -> list[dict]:
    """Fallback: if a docx has no headings, split by paragraph groups."""
    doc = Document(filepath)
    stem = filepath.stem
    parts = stem.split("_")
    chapter = parts[-1] if len(parts) > 1 else stem

    all_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)

    if not all_text:
        return []

    full = "\n".join(all_text)
    if len(full) < MIN_CHUNK_SIZE:
        return []

    # Split into chunks of ~MAX_CHUNK_SIZE characters
    chunks = []
    current = []
    current_len = 0
    chunk_idx = 0

    for line in all_text:
        current.append(line)
        current_len += len(line)
        if current_len >= MAX_CHUNK_SIZE:
            chunk_idx += 1
            chunks.append({
                "path": f"{chapter} / Part {chunk_idx}",
                "content": "\n".join(current),
                "source_file": filepath.name,
            })
            current = []
            current_len = 0

    if current:
        chunk_idx += 1
        chunks.append({
            "path": f"{chapter} / Part {chunk_idx}",
            "content": "\n".join(current),
            "source_file": filepath.name,
        })

    return chunks


def add_keywords_with_llm(chunks: list[dict], llm: LLMClient) -> list[dict]:
    """Use Claude Haiku to extract keywords for each chunk."""
    total = len(chunks)
    for i, chunk in enumerate(chunks):
        content = chunk["content"]
        # Truncate very long chunks for keyword extraction
        if len(content) > MAX_CHUNK_SIZE:
            content = content[:MAX_CHUNK_SIZE] + "\n...(truncated)"

        prompt = KEYWORD_PROMPT.format(
            heading=chunk["path"],
            content=content,
        )

        logger.info(f"  [{i+1}/{total}] Extracting keywords for: {chunk['path']}")
        result = llm.call_json(prompt, model_key="haiku", max_tokens=500)

        if isinstance(result, list):
            chunk["keywords"] = result
        else:
            # Fallback: extract simple keywords from content
            logger.warning(f"    LLM keyword extraction failed, using fallback")
            chunk["keywords"] = _fallback_keywords(chunk)

    return chunks


def _fallback_keywords(chunk: dict) -> list[str]:
    """Simple regex-based keyword extraction as fallback."""
    content = chunk["content"]
    # Find English medical abbreviations (2-6 uppercase letters)
    abbrevs = set(re.findall(r'\b[A-Z]{2,6}\b', content))
    # Find English medical terms (capitalized words 4+ chars)
    terms = set(re.findall(r'\b[A-Z][a-z]{3,}\b', content))
    keywords = list(abbrevs | terms)[:10]
    return keywords


def main():
    parser = argparse.ArgumentParser(description="Process work manuals into chunks")
    parser.add_argument("--dept", default="newborn", help="Department name (default: newborn)")
    parser.add_argument("--no-keywords", action="store_true", help="Skip LLM keyword extraction")
    parser.add_argument("--manuals-dir", default=None, help="Override manuals directory path")
    args = parser.parse_args()

    # Paths
    project_root = Path(__file__).resolve().parent.parent
    manuals_dir = Path(args.manuals_dir) if args.manuals_dir else project_root / "config" / "manuals"
    output_path = project_root / "data" / args.dept / "manual_chunks.json"

    if not manuals_dir.exists():
        logger.error(f"Manuals directory not found: {manuals_dir}")
        sys.exit(1)

    docx_files = sorted(manuals_dir.glob("*.docx"))
    if not docx_files:
        logger.error(f"No .docx files found in {manuals_dir}")
        sys.exit(1)

    logger.info(f"Found {len(docx_files)} .docx files in {manuals_dir}")

    # Extract chunks from all files
    all_chunks = []
    for filepath in docx_files:
        logger.info(f"\nProcessing: {filepath.name}")
        chunks = extract_chunks_from_docx(filepath)

        # If no heading-based chunks, use size-based splitting
        if not chunks:
            logger.info(f"  No headings found, splitting by size...")
            chunks = extract_chunks_from_docx_by_size(filepath)

        logger.info(f"  -> {len(chunks)} chunks")
        all_chunks.extend(chunks)

    logger.info(f"\nTotal chunks: {len(all_chunks)}")

    # Extract keywords
    if not args.no_keywords:
        logger.info("\nExtracting keywords with Claude Haiku...")
        try:
            llm = LLMClient()
            all_chunks = add_keywords_with_llm(all_chunks, llm)
            usage = llm.get_usage_summary()
            logger.info(f"\nLLM usage: {usage['total_calls']} calls, ~${usage['estimated_cost_usd']:.4f}")
        except ValueError as e:
            logger.warning(f"\n{e}")
            logger.warning("Falling back to regex-based keyword extraction...")
            for chunk in all_chunks:
                chunk["keywords"] = _fallback_keywords(chunk)
    else:
        logger.info("\nSkipping LLM keywords (--no-keywords), using regex fallback...")
        for chunk in all_chunks:
            chunk["keywords"] = _fallback_keywords(chunk)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(all_chunks, f, ensure_ascii=False, indent=2)

    logger.info(f"\nSaved {len(all_chunks)} chunks to {output_path}")


if __name__ == "__main__":
    main()
